from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "work_assets"
CHARTS = ASSETS / "charts"
CHARTS.mkdir(parents=True, exist_ok=True)
OUT = ROOT / "我的世界模组视频赛道深度可视化报告.docx"
SOURCE_IMAGE = ASSETS / "source_unzip" / "word" / "media" / "image1.jpeg"


COLORS = {
    "navy": "0B2545",
    "blue": "2E74B5",
    "cyan": "2CA6A4",
    "gold": "C18C2B",
    "red": "B54747",
    "green": "3A7D44",
    "gray": "5A6777",
    "light": "F4F7FB",
    "line": "D9E2EC",
    "ink": "1F2937",
    "white": "FFFFFF",
}


def pick_chinese_font():
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for item in candidates:
        if Path(item).exists():
            return item
    return None


FONT_PATH = pick_chinese_font()


def font(size, bold=False):
    if FONT_PATH:
        try:
            return ImageFont.truetype(FONT_PATH, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


def draw_center(draw, xy, text, fnt, fill, anchor="mm"):
    draw.text(xy, text, font=fnt, fill=fill, anchor=anchor)


def hex_to_rgb(value):
    value = value.lstrip("#")
    return tuple(int(value[i:i + 2], 16) for i in (0, 2, 4))


def save_image(img, path):
    img.save(path, quality=95)
    return path


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=COLORS["ink"], size=9.2, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell)


def set_table_borders(table, color="D9E2EC", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def add_paragraph(doc, text="", size=10.5, color=COLORS["ink"], bold=False, italic=False,
                  align=None, before=0, after=6, line=1.1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_callout(doc, label, body, fill="F4F7FB", accent=COLORS["blue"]):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    set_table_borders(table, color="E5EDF5", size="4")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    set_run_font(r, size=9.5, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2, color=COLORS["ink"])


def add_metric_strip(doc):
    data = [
        ("恐怖模组", "起量快", "爆款依赖强"),
        ("机械动力", "长尾强", "制作周期长"),
        ("推荐打法", "双线并行", "恐怖拉新 + 机械沉淀"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_widths(table, [2.1, 2.1, 2.3])
    set_table_borders(table, color="E2E8F0", size="4")
    for i, (a, b, c) in enumerate(data):
        cell = table.cell(0, i)
        shade_cell(cell, ["EEF6FF", "F0FDF4", "FFF7ED"][i])
        cell_margins(cell, top=140, bottom=140, start=140, end=140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(a)
        set_run_font(r, size=10.5, color=[COLORS["blue"], COLORS["green"], COLORS["gold"]][i], bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(b)
        set_run_font(r2, size=15, color=COLORS["navy"], bold=True)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(c)
        set_run_font(r3, size=8.5, color=COLORS["gray"])


def save_bar_chart():
    labels = ["起量速度", "完播/互动", "制作门槛", "收入稳定性", "资产沉淀", "商业延展"]
    horror = [9, 9, 5, 4, 3, 6]
    create = [4, 6, 8, 8, 9, 8]
    img = Image.new("RGB", (1400, 720), "white")
    draw = ImageDraw.Draw(img)
    draw_center(draw, (700, 58), "两条赛道的能力画像对比", font(34), hex_to_rgb(COLORS["navy"]))
    draw.rounded_rectangle((470, 92, 930, 140), radius=18, fill="#F8FAFC", outline="#E2E8F0")
    draw.rectangle((505, 112, 535, 130), fill="#B54747")
    draw.text((545, 104), "恐怖模组", font=font(24), fill="#334155")
    draw.rectangle((690, 112, 720, 130), fill="#2E74B5")
    draw.text((730, 104), "机械动力", font=font(24), fill="#334155")
    left, top, width, height = 115, 175, 1190, 420
    for i in range(6):
        y = top + height - i * height / 5
        draw.line((left, y, left + width, y), fill="#E2E8F0", width=2)
        draw.text((55, y - 13), str(i * 2), font=font(20), fill="#64748B")
    group_w = width / len(labels)
    bar_w = 34
    for i, label in enumerate(labels):
        cx = left + group_w * i + group_w / 2
        for offset, val, color in [(-22, horror[i], "#B54747"), (22, create[i], "#2E74B5")]:
            bh = val / 10 * height
            x0 = cx + offset - bar_w / 2
            y0 = top + height - bh
            draw.rounded_rectangle((x0, y0, x0 + bar_w, top + height), radius=8, fill=color)
            draw.text((x0 + 5, y0 - 30), str(val), font=font(18), fill="#334155")
        draw_center(draw, (cx, top + height + 45), label, font(21), "#334155")
    draw.text((42, 165), "相对评分（10分制）", font=font(20), fill="#64748B")
    path = CHARTS / "score_bar.png"
    return save_image(img, path)


def save_radar_chart():
    labels = ["内容爆发", "情绪传播", "搜索长尾", "付费资产", "制作复杂度", "品牌沉淀"]
    horror = [9, 9, 3, 3, 5, 5]
    create = [4, 5, 9, 9, 8, 8]
    img = Image.new("RGB", (1000, 920), "white")
    draw = ImageDraw.Draw(img, "RGBA")
    cx, cy, max_r = 500, 485, 280
    draw_center(draw, (500, 58), "内容价值结构雷达图", font(34), hex_to_rgb(COLORS["navy"]))
    for level in [2, 4, 6, 8, 10]:
        pts = []
        for i in range(len(labels)):
            ang = -math.pi / 2 + i * 2 * math.pi / len(labels)
            r = max_r * level / 10
            pts.append((cx + r * math.cos(ang), cy + r * math.sin(ang)))
        draw.line(pts + [pts[0]], fill="#CBD5E1", width=2)
    for i, label in enumerate(labels):
        ang = -math.pi / 2 + i * 2 * math.pi / len(labels)
        draw.line((cx, cy, cx + max_r * math.cos(ang), cy + max_r * math.sin(ang)), fill="#E2E8F0", width=2)
        tx = cx + (max_r + 58) * math.cos(ang)
        ty = cy + (max_r + 58) * math.sin(ang)
        draw_center(draw, (tx, ty), label, font(24), "#334155")

    def poly(values, color, alpha):
        pts = []
        for i, val in enumerate(values):
            ang = -math.pi / 2 + i * 2 * math.pi / len(values)
            r = max_r * val / 10
            pts.append((cx + r * math.cos(ang), cy + r * math.sin(ang)))
        draw.polygon(pts, fill=hex_to_rgb(color) + (alpha,))
        draw.line(pts + [pts[0]], fill=color, width=5)

    poly(horror, "#B54747", 42)
    poly(create, "#2E74B5", 42)
    draw.rectangle((325, 105, 355, 123), fill="#B54747")
    draw.text((365, 96), "恐怖模组", font=font(24), fill="#334155")
    draw.rectangle((520, 105, 550, 123), fill="#2E74B5")
    draw.text((560, 96), "机械动力", font=font(24), fill="#334155")
    path = CHARTS / "radar.png"
    return save_image(img, path)


def save_timeline_chart():
    weeks = ["第1周", "第2周", "第3周", "第4周", "第5-6周", "第7-8周"]
    horror = [8, 9, 7, 6, 5, 4]
    create = [2, 3, 4, 6, 7, 8]
    img = Image.new("RGB", (1400, 650), "white")
    draw = ImageDraw.Draw(img)
    draw_center(draw, (700, 58), "8周内容组合节奏建议", font(34), hex_to_rgb(COLORS["navy"]))
    left, top, width, height = 125, 155, 1160, 360
    for i in range(6):
        y = top + height - i * height / 5
        draw.line((left, y, left + width, y), fill="#E2E8F0", width=2)
    def points(values):
        return [(left + i * width / (len(values)-1), top + height - v / 10 * height) for i, v in enumerate(values)]
    hp, cp = points(horror), points(create)
    draw.line(hp, fill="#B54747", width=6, joint="curve")
    draw.line(cp, fill="#2E74B5", width=6, joint="curve")
    for pts, color in [(hp, "#B54747"), (cp, "#2E74B5")]:
        for x, y in pts:
            draw.ellipse((x - 11, y - 11, x + 11, y + 11), fill=color, outline="white", width=4)
    for i, week in enumerate(weeks):
        x = left + i * width / (len(weeks)-1)
        draw_center(draw, (x, top + height + 42), week, font(22), "#334155")
    draw.rectangle((410, 100, 440, 118), fill="#B54747")
    draw.text((450, 92), "恐怖模组：拉新热度", font=font(23), fill="#334155")
    draw.rectangle((720, 100, 750, 118), fill="#2E74B5")
    draw.text((760, 92), "机械动力：沉淀价值", font=font(23), fill="#334155")
    draw.text((40, 145), "经营强度/收益贡献", font=font(20), fill="#64748B")
    path = CHARTS / "timeline.png"
    return save_image(img, path)


def save_funnel_chart():
    img = Image.new("RGB", (1400, 620), "white")
    draw = ImageDraw.Draw(img)
    steps = [
        ("短视频爆点", "恐怖反应/名场面"),
        ("主页沉淀", "系列合集/置顶教程"),
        ("私域承接", "社群/直播/会员"),
        ("资产销售", "蓝图/存档/整合包"),
    ]
    fills = ["#FEE2E2", "#DBEAFE", "#E0F2FE", "#DCFCE7"]
    edge = ["#B54747", "#2E74B5", "#0284C7", "#3A7D44"]
    draw_center(draw, (700, 78), "从爆款流量到长期收入的承接链路", font(34), hex_to_rgb(COLORS["navy"]))
    x0, y0, w, h, gap = 75, 230, 260, 160, 65
    for i, (title, desc) in enumerate(steps):
        x = x0 + i * (w + gap)
        draw.rounded_rectangle((x, y0, x + w, y0 + h), radius=26, fill=fills[i], outline=edge[i], width=4)
        draw_center(draw, (x + w / 2, y0 + 62), title, font(29), "#0B2545")
        draw_center(draw, (x + w / 2, y0 + 108), desc, font(22), "#465366")
        if i < len(steps) - 1:
            ax0 = x + w + 12
            ax1 = x + w + gap - 12
            ay = y0 + h / 2
            draw.line((ax0, ay, ax1, ay), fill="#64748B", width=5)
            draw.polygon([(ax1, ay), (ax1 - 18, ay - 12), (ax1 - 18, ay + 12)], fill="#64748B")
    path = CHARTS / "funnel.png"
    return save_image(img, path)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
        s.header_distance = Inches(0.49)
        s.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    return doc


def add_cover(doc):
    add_paragraph(doc, "赛道分析报告", size=10, color=COLORS["gold"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    add_paragraph(doc, "《我的世界》模组视频赛道", size=25, color=COLORS["navy"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_paragraph(doc, "恐怖模组与机械动力的内容增长、变现路径与运营打法", size=13, color=COLORS["gray"], align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    add_metric_strip(doc)
    add_paragraph(doc, "", after=12)
    add_callout(
        doc,
        "核心结论",
        "恐怖模组更适合做流量破圈和短期爆发，机械动力更适合做教程长尾、蓝图/存档售卖和创作者品牌沉淀。最优策略不是二选一，而是用恐怖内容吸引新观众，用机械动力内容承接信任和复购。",
        fill="F8FAFC",
        accent=COLORS["navy"],
    )
    table = doc.add_table(rows=1, cols=3)
    set_table_widths(table, [2.1, 2.1, 2.3])
    set_table_borders(table, color="E2E8F0", size="4")
    highlights = [
        ("报告重点", "赛道对比、案例拆解、变现路径"),
        ("关键图表", "能力评分、价值雷达、承接链路"),
        ("落地输出", "8周路线图、风险控制、最终建议"),
    ]
    for i, (title, body) in enumerate(highlights):
        cell = table.cell(0, i)
        shade_cell(cell, "FFFFFF")
        set_cell_text(cell, title + "\n" + body, bold=False, color=COLORS["ink"], size=8.9, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "报告口径：本报告基于原始简报中给出的案例、粉丝量、点赞量、收入锚点与赛道判断重写扩展；未额外进行第三方数据核验，商业测算采用相对评分与运营假设。", size=8.5, color=COLORS["gray"], italic=True, after=10)


def add_comparison_table(doc):
    add_heading(doc, "1. 赛道总览：快钱与长钱的根本差异", 1)
    add_paragraph(doc, "两类内容的核心差别，不在于谁更“好看”，而在于观众为什么停留、为什么关注、以及为什么付费。恐怖模组卖的是即时情绪，机械动力卖的是知识、方案和可复用资产。")
    rows = [
        ("核心驱动力", "惊吓、反应、悬念、剧情反转", "建造、自动化、教程、工程奇观"),
        ("内容生命周期", "短视频爆发强，热点衰减快", "搜索长尾明显，教程可持续获流量"),
        ("典型变现", "流量分成、发行人计划、直播打赏、商单", "蓝图/存档售卖、整合包分发、教程长尾、会员订阅"),
        ("创作者要求", "表演力、更新速度、情绪感染力", "系统理解、工程拆解、耐心制作、表达清晰"),
        ("主要风险", "审美疲劳、题材同质化、收入波动", "起量慢、周期长、学习门槛高"),
        ("建议定位", "拉新入口、阶段性爆点、直播互动素材", "品牌资产、稳定收入、粉丝信任承接"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_widths(table, [1.35, 2.55, 2.6])
    set_table_borders(table)
    headers = ["比较维度", "恐怖模组", "机械动力"]
    for i, h in enumerate(headers):
        shade_cell(table.cell(0, i), "E8EEF5")
        set_cell_text(table.cell(0, i), h, bold=True, color=COLORS["navy"], size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for label, horror, create in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=COLORS["navy"], size=9.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], horror, size=9.0)
        set_cell_text(cells[2], create, size=9.0)


def add_charts(doc, bar, radar):
    add_heading(doc, "2. 可视化判断：谁负责爆发，谁负责沉淀", 1)
    add_paragraph(doc, "从相对评分看，恐怖模组在起量速度、完播互动上显著占优；机械动力在收入稳定性、资产沉淀与商业延展上更强。")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(bar), width=Inches(6.2))
    add_paragraph(doc, "图1：两类内容的相对能力评分。评分用于经营决策比较，不代表平台官方数据。", size=8.5, color=COLORS["gray"], align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(radar), width=Inches(4.7))
    add_paragraph(doc, "图2：内容价值结构雷达图。恐怖内容强在传播，机械动力强在复用与转化。", size=8.5, color=COLORS["gray"], align=WD_ALIGN_PARAGRAPH.CENTER)


def add_track_sections(doc):
    add_heading(doc, "3. 恐怖模组赛道：情绪驱动的爆款机器", 1)
    add_callout(doc, "一句话判断", "恐怖模组适合先把账号做“热”：用高情绪、强反应、短周期选题快速积累播放、互动和关注。", fill="FEF2F2", accent=COLORS["red"])
    add_paragraph(doc, "恐怖模组的内容机制非常直接：观众期待创作者被吓到、反应失控、剧情反转，视频天然具备分享动机。它不要求观众理解复杂系统，只要镜头节奏、音效、标题和反应足够强，就容易形成完播。")
    for text in [
        "选题核心：围绕“诡月”“未知生物”“极限生存”“多人联机整活”等关键词做强悬念包装。",
        "变现核心：优先承接发行人计划、直播打赏和爆款流量分成，商单适合在账号有稳定互动后接入。",
        "运营核心：保持模组新鲜度，避免同一种惊吓套路反复出现；用系列化标题降低观众理解成本。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)
    add_paragraph(doc, "原文数据锚点：Mine 威杰抖音账号约 99 万粉，“诡月恐怖生存”系列获高赞；单条爆款恐怖模组视频可出现 17 万+点赞量。", size=9.2, color=COLORS["gray"], italic=True)
    if SOURCE_IMAGE.exists():
        p = add_paragraph(doc, "案例截图：Mine 威杰账号主页与恐怖内容矩阵", size=8.8, color=COLORS["gray"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        p.paragraph_format.keep_with_next = True
        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.add_run().add_picture(str(SOURCE_IMAGE), width=Inches(1.7))

    add_heading(doc, "4. 机械动力赛道：创造驱动的长期资产", 1)
    add_callout(doc, "一句话判断", "机械动力适合把账号做“厚”：用教程、工程展示、蓝图和整合包沉淀可搜索、可复购、可长期分发的内容资产。", fill="EFF6FF", accent=COLORS["blue"])
    add_paragraph(doc, "机械动力的价值不只来自播放量，还来自“能不能解决问题”。观众常常带着明确需求搜索：如何搭建某条产线、如何理解齿轮/传送带/动力系统、如何下载可复用蓝图。只要内容足够清晰，就可能长期带来播放和付费。")
    for text in [
        "选题核心：从“新手能复刻的小工程”到“高阶自动化奇观”分层设计，避免一开始就过度复杂。",
        "变现核心：将视频中的工程沉淀为蓝图、存档、整合包、会员教程与答疑社群。",
        "运营核心：用清晰目录、章节、材料清单和下载入口提高转化，降低观众复刻成本。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)
    add_paragraph(doc, "原文数据锚点：机械动力教程具备数月乃至数年长尾；地图/存档可售卖形成被动收入；模组开发者和获奖模组存在月入 2 万至 3 万元级别案例。", size=9.2, color=COLORS["gray"], italic=True)


def add_strategy(doc, funnel, timeline):
    add_heading(doc, "5. 推荐策略：恐怖拉新，机械动力变现", 1)
    add_paragraph(doc, "如果目标是做一个能赚钱且不容易断档的账号，建议不要把两条路线割裂。更稳的打法是“前端情绪内容拉新，后端技术内容承接”。")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(funnel), width=Inches(6.3))
    add_paragraph(doc, "图3：从爆款流量到长期收入的承接链路。", size=8.5, color=COLORS["gray"], align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(timeline), width=Inches(6.2))
    add_paragraph(doc, "图4：8周内容组合节奏建议。前期用恐怖内容拿曝光，中后期提高机械动力内容占比。", size=8.5, color=COLORS["gray"], align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()
    add_heading(doc, "6. 8周执行路线图", 1)
    rows = [
        ("第1-2周", "建立人设与爆点", "恐怖模组短视频 6-8 条；直播 1-2 场", "验证标题、封面、反应强度"),
        ("第3-4周", "加入机械动力承接", "机械动力入门教程 3 条；置顶合集", "观察收藏、完播、私信需求"),
        ("第5-6周", "做可售资产雏形", "蓝图/存档小样；直播演示工程", "测试下载意愿与付费价格"),
        ("第7-8周", "形成双线栏目", "恐怖系列维持热度；机械系列稳定更新", "沉淀社群、会员、整合包分发"),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_widths(table, [1.05, 1.55, 2.45, 1.45])
    set_table_borders(table)
    for i, h in enumerate(["阶段", "目标", "内容动作", "关键指标"]):
        shade_cell(table.cell(0, i), "E8EEF5")
        set_cell_text(table.cell(0, i), h, bold=True, color=COLORS["navy"], size=9.3, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, bold=(i == 0), color=COLORS["navy"] if i == 0 else COLORS["ink"], size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER if i in [0, 1, 3] else None)


def add_monetization_and_risks(doc):
    add_heading(doc, "7. 变现路径与风险控制", 1)
    rows = [
        ("流量分成", "恐怖爆款、机械教程均可", "恐怖短期更强；机械长尾更稳", "关注完播率、搜索流量、更新频率"),
        ("发行人计划", "恐怖/剧情/整活视频", "强依赖平台活动与转化", "避免标题党透支信任"),
        ("直播打赏", "恐怖联机、挑战、观众点题", "互动强但消耗人设", "控制直播频次，保留高光切片"),
        ("蓝图/存档售卖", "机械动力、地图解谜、整合包", "可形成复用资产", "提供预览、说明、售后与版本记录"),
        ("会员/社群", "高粘性粉丝、教程深度需求", "稳定但需持续交付", "明确权益，避免只卖情绪"),
        ("商单合作", "账号有稳定互动后", "单次收入高但破坏体验风险", "只接符合 MC 场景的产品/服务"),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_widths(table, [1.25, 1.65, 1.75, 1.85])
    set_table_borders(table)
    for i, h in enumerate(["变现方式", "适配场景", "收益特点", "风险控制"]):
        shade_cell(table.cell(0, i), "F2F4F7")
        set_cell_text(table.cell(0, i), h, bold=True, color=COLORS["navy"], size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, bold=(i == 0), color=COLORS["navy"] if i == 0 else COLORS["ink"], size=8.45, align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else None)

    doc.add_page_break()
    add_heading(doc, "8. 最终建议", 1)
    add_callout(
        doc,
        "主策略",
        "以恐怖模组作为流量入口，占前期内容的 60%-70%；同时从第 3 周开始持续发布机械动力教程和工程展示。等账号拥有稳定关注后，把机械动力内容升级成蓝图、存档、整合包和会员服务，逐步把收入从爆款依赖转为资产复利。",
        fill="F0FDF4",
        accent=COLORS["green"],
    )
    for text in [
        "如果创作者更会表演：先主打恐怖，但必须设计机械动力或地图资产作为后端。",
        "如果创作者更懂技术：先做机械动力，但每周加入 1-2 条恐怖/挑战向短视频，提高破圈概率。",
        "如果团队只有一个人：优先做“恐怖短视频 + 机械小教程”，不要一开始就做大型整合包。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)

    add_heading(doc, "9. 可直接开拍的栏目选题库", 1)
    rows = [
        ("恐怖拉新", "《我在诡月里活过第 7 天》", "强标题、强反应、短节奏，适合切片和直播预告"),
        ("恐怖互动", "观众投票决定下一晚规则", "把评论区变成选题池，提高互动与复访"),
        ("机械入门", "10 分钟搭好第一条自动矿物产线", "降低门槛，承接新粉的学习需求"),
        ("机械资产", "本期工程蓝图/存档开放下载", "把观看兴趣转成可复用资产与付费入口"),
        ("双线联动", "用机械动力基地对抗恐怖事件", "把爆点剧情和工程能力合在同一世界观里"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_widths(table, [1.35, 2.35, 2.8])
    set_table_borders(table)
    for i, h in enumerate(["栏目方向", "标题样例", "运营目的"]):
        shade_cell(table.cell(0, i), "E8EEF5")
        set_cell_text(table.cell(0, i), h, bold=True, color=COLORS["navy"], size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, bold=(i == 0), color=COLORS["navy"] if i == 0 else COLORS["ink"], size=8.6, align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else None)


def build():
    bar = save_bar_chart()
    radar = save_radar_chart()
    timeline = save_timeline_chart()
    funnel = save_funnel_chart()

    doc = setup_document()
    add_cover(doc)
    doc.add_page_break()
    add_comparison_table(doc)
    add_charts(doc, bar, radar)
    add_track_sections(doc)
    add_strategy(doc, funnel, timeline)
    add_monetization_and_risks(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("《我的世界》模组视频赛道深度可视化报告")
    set_run_font(run, size=8.5, color=COLORS["gray"])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
